#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 USER_MANUAL.md 转换为带封面、目录、统一样式（但无插图）的 Word 文档。
依赖：python-docx
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, 'USER_MANUAL.md')
OUT = os.path.join(ROOT, 'USER_MANUAL.docx')

# ---------- 颜色定义 ----------
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
INDIGO_DARK = RGBColor(0x37, 0x30, 0xA3)
SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
SLATE_400 = RGBColor(0x94, 0xA3, 0xB8)
EMERALD = RGBColor(0x05, 0x96, 0x69)
ROSE = RGBColor(0xE1, 0x1D, 0x48)
AMBER = RGBColor(0xD9, 0x77, 0x06)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CN_FONT = '微软雅黑'
EN_FONT = 'Calibri'


def set_cell_shading(cell, color_hex):
    """给表格单元格设置背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex='E2E8F0', size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement('w:' + edge)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_runs_with_inline(paragraph, text):
    """
    解析简单的行内格式：**bold**、`code`、普通文本。
    """
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = EN_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)
        token = m.group(0)
        if token.startswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = SLATE_900
        elif token.startswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10.5)
            run.font.color.rgb = INDIGO
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            rfonts.set(qn('w:eastAsia'), CN_FONT)
        run.font.name = EN_FONT
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = EN_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), CN_FONT)


def set_paragraph_cjk_font(paragraph, size=11, color=SLATE_700, bold=False):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.font.name = EN_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), CN_FONT)


def style_heading(paragraph, level):
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    colors = {1: INDIGO_DARK, 2: INDIGO, 3: SLATE_900, 4: SLATE_900}
    for run in paragraph.runs:
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = colors.get(level, SLATE_900)
        run.bold = True
        run.font.name = EN_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), CN_FONT)
    pf = paragraph.paragraph_format
    if level == 1:
        pf.space_before = Pt(28)
        pf.space_after = Pt(14)
        # 一级标题下方加分隔线
        add_bottom_border(paragraph, color='E0E7FF', size=12)
    elif level == 2:
        pf.space_before = Pt(20)
        pf.space_after = Pt(10)
    else:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)


def add_bottom_border(paragraph, color='E0E7FF', size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_shaded_paragraph(doc, text, fill_hex, border_hex, text_color, icon=''):
    """
    添加一个带背景色和左边框的段落（用于 > 提示、警告框等）。
    实现方式：用 1x1 表格模拟
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill_hex)
    # 左边粗边、其他细边
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'right', 'bottom'):
        b = OxmlElement('w:' + edge)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), border_hex)
        tc_borders.append(b)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), border_hex)
    tc_borders.append(left)
    tc_pr.append(tc_borders)
    # 内容
    cell.text = ''
    p = cell.paragraphs[0]
    full = (icon + ' ' if icon else '') + text
    add_runs_with_inline(p, full)
    set_paragraph_cjk_font(p, size=10.5, color=text_color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # 表格宽度铺满
    set_table_width(table, Cm(16))
    # 加点间距
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def set_table_width(table, width):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(int(width.twips)))


# ---------- 文档构建 ----------
def build_cover(doc):
    # 顶部留白
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # 标签
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('USER GUIDE')
    r.font.size = Pt(11)
    r.font.color.rgb = INDIGO
    r.bold = True
    r.font.name = EN_FONT
    p.paragraph_format.space_after = Pt(18)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Ebbinghaus Memo')
    r.font.size = Pt(42)
    r.font.color.rgb = SLATE_900
    r.bold = True
    r.font.name = EN_FONT

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('用户操作手册')
    r.font.size = Pt(22)
    r.font.color.rgb = INDIGO
    r.bold = True
    r.font.name = EN_FONT
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    rpr.append(rfonts)
    p.paragraph_format.space_after = Pt(24)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('————  ————  ————')
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE_400
    p.paragraph_format.space_after = Pt(24)

    # 简介
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('艾宾浩斯智能记忆引擎 · 多语言背单词 Web 应用')
    r.font.size = Pt(13)
    r.font.color.rgb = SLATE_500
    r.font.name = EN_FONT
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    rpr.append(rfonts)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('从注册账号到掌握单词 · 完整功能指南')
    r.font.size = Pt(11)
    r.font.color.rgb = SLATE_400
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    rpr.append(rfonts)

    # 底部留白
    for _ in range(10):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # 版本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Version 1.0  ·  2026.07')
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE_400
    r.font.name = EN_FONT

    doc.add_page_break()


def build_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run('目录')
    style_heading(p, 1)

    items = [
        ('1. 产品简介', '产品定位、核心特性'),
        ('2. 账号注册与登录', '注册、登录、忘记密码处理'),
        ('3. 主界面导航', '四个主视图、语言切换'),
        ('4. 仪表盘（首页）', '统计卡、趋势图、阶段分布'),
        ('5. 词库管理', '添加、批量导入、筛选、详情'),
        ('6. 复习会话', '闪卡 / 拼写、错词循环'),
        ('7. 个人资料与勋章墙', '改资料、改密码、12 枚勋章'),
        ('8. 记忆算法原理', '6 阶段间隔、奖惩规则'),
        ('9. 管理员专属功能', '时光机、数据库重置'),
        ('10. 常见问题（FAQ）', '8 条常见问题答疑'),
    ]
    for title, desc in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run('▪ ')
        r1.font.color.rgb = INDIGO
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(title)
        r2.bold = True
        r2.font.size = Pt(11.5)
        r2.font.color.rgb = SLATE_900
        r2.font.name = EN_FONT
        rpr = r2._element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), CN_FONT)
        rpr.append(rfonts)
        r3 = p.add_run('  — ' + desc)
        r3.font.size = Pt(10)
        r3.font.color.rgb = SLATE_500
        r3.font.name = EN_FONT
        rpr = r3._element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), CN_FONT)
        rpr.append(rfonts)

    doc.add_page_break()


def add_table_from_rows(doc, header, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        run.font.name = EN_FONT
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), CN_FONT)
        rpr.append(rfonts)
        set_cell_shading(cell, '4F46E5')
        set_cell_borders(cell, color_hex='3730A3', size=4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs_with_inline(p, str(val))
            set_paragraph_cjk_font(p, size=10.5, color=SLATE_700)
            # 隔行底色
            if ri % 2 == 1:
                set_cell_shading(cell, 'F8FAFC')
            set_cell_borders(cell, color_hex='E2E8F0', size=4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 列宽
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    # 整表宽度
    total = sum(col_widths_cm) if col_widths_cm else 16
    set_table_width(table, Cm(total))
    # 空一行
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_para(doc, text, size=11, color=SLATE_700, bold=False, indent_cm=0, space_after=6):
    p = doc.add_paragraph()
    add_runs_with_inline(p, text)
    set_paragraph_cjk_font(p, size=size, color=color, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    add_runs_with_inline(p, text)
    set_paragraph_cjk_font(p, size=11, color=SLATE_700)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    add_runs_with_inline(p, text)
    set_paragraph_cjk_font(p, size=11, color=SLATE_700)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4


def add_code_block(doc, text):
    """用浅灰底框模拟代码块"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F1F5F9')
    set_cell_borders(cell, color_hex='E2E8F0', size=4)
    cell.text = ''
    for i, line in enumerate(text.split('\n')):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
        r.font.color.rgb = INDIGO_DARK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
    set_table_width(table, Cm(16))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_tip(doc, text, kind='tip'):
    styles = {
        'tip':    ('ECFDF5', '10B981', SLATE_700, '💡'),
        'warn':   ('FFFBEB', 'F59E0B', SLATE_700, '⚠️'),
        'danger': ('FEF2F2', 'EF4444', SLATE_700, '⚠️'),
        'info':   ('EFF6FF', '3B82F6', SLATE_700, 'ℹ️'),
    }
    fill, border, color, icon = styles[kind]
    add_shaded_paragraph(doc, text, fill, border, color, icon=icon)


def build_body(doc):
    # ====================== 1. 产品简介 ======================
    p = doc.add_paragraph(); p.add_run('1. 产品简介'); style_heading(p, 1)

    add_para(doc, '**Ebbinghaus Memo** 是一款基于艾宾浩斯遗忘曲线的多语言背单词应用。')
    p = doc.add_paragraph(); p.add_run('核心特性：'); style_heading(p, 3)
    for t in [
        '**6 阶段间隔重复算法（SRS）**：按照 [1, 2, 4, 7, 15, 30] 天的科学节奏自动排期复习',
        '**AI 自动生成释义**：添加新词时由 AI 自动生成中文释义、双语例句、艾宾浩斯助记',
        '**5 种目标语言词库**：英语 / 日语 / 西班牙语 / 法语 / 葡萄牙语',
        '**6 种 UI 界面语言**：中文 + 5 种目标语言任意切换',
        '**错词循环重考**：复习中答错的词会在同一次会话内立即重考，直到记住',
        '**勋章成就系统**：12 枚勋章覆盖"学习坚持 / 词库探索 / 终极掌握"三大维度',
        '**真人发音 + 浏览器 TTS 兜底**：每个单词都能听到标准发音',
    ]:
        add_bullet(doc, t)

    p = doc.add_paragraph(); p.add_run('开始使用'); style_heading(p, 3)
    add_para(doc, '打开浏览器访问应用地址（本地开发环境默认为 `http://localhost:3003`），注册账号即可开始。')

    doc.add_page_break()

    # ====================== 2. 账号注册与登录 ======================
    p = doc.add_paragraph(); p.add_run('2. 账号注册与登录'); style_heading(p, 1)

    p = doc.add_paragraph(); p.add_run('2.1 注册新账号'); style_heading(p, 2)
    add_numbered(doc, '在登录页顶部点击 **"邮箱注册"** Tab。')
    add_numbered(doc, '填写以下信息：')
    add_table_from_rows(
        doc,
        ['字段', '说明', '约束'],
        [
            ['电子邮箱', '账号唯一标识，登录时使用', '必填，邮箱格式'],
            ['密码', '登录凭证', '至少 6 位'],
            ['个性昵称', '导航栏 / 个人资料显示名', '必填'],
            ['每日复习目标（词）', '系统统计与提醒依据', '5–200，默认 15'],
        ],
        col_widths_cm=[4, 7, 5]
    )
    add_numbered(doc, '点击 **"完成注册并生成词库"**。')
    add_numbered(doc, '注册成功后自动登录并跳转到仪表盘。')
    add_tip(doc, '关于"每日复习目标"：用于系统统计与提醒。建议初学者保持默认 15，熟练后可到个人资料页修改。', kind='tip')

    p = doc.add_paragraph(); p.add_run('2.2 登录已有账号'); style_heading(p, 2)
    add_numbered(doc, '在登录页保持 **"登录账号"** Tab。')
    add_numbered(doc, '输入邮箱与密码。')
    add_numbered(doc, '点击 **"登录智能词库"**。')

    p = doc.add_paragraph(); p.add_run('2.3 忘记密码？'); style_heading(p, 2)
    add_tip(doc, '应用不提供邮箱找回密码功能。如忘记密码，请联系管理员重置，或通过已登录的其他设备到个人资料页修改。', kind='warn')

    doc.add_page_break()

    # ====================== 3. 主界面导航 ======================
    p = doc.add_paragraph(); p.add_run('3. 主界面导航'); style_heading(p, 1)
    add_para(doc, '登录成功后，顶部会出现常驻导航栏（在手机端会折叠为图标）。')

    p = doc.add_paragraph(); p.add_run('3.1 四个主视图'); style_heading(p, 2)
    add_table_from_rows(
        doc,
        ['导航项', '图标', '功能说明'],
        [
            ['仪表盘', 'LayoutDashboard', '查看统计、开始今日复习、查看趋势图'],
            ['复习', 'BookOpen', '进入复习会话（有待复习词时，右上角显示红色数字徽章）'],
            ['词库', 'GraduationCap', '管理所有单词：增删改查、批量导入、筛选搜索'],
            ['个人（显示昵称）', 'User', '修改资料、改密码、查看勋章墙、退出登录'],
        ],
        col_widths_cm=[4, 3.5, 8.5]
    )
    add_para(doc, '点击左上角 Logo 也可回到仪表盘。')

    p = doc.add_paragraph(); p.add_run('3.2 语言切换（右上角两个下拉）'); style_heading(p, 2)
    p = doc.add_paragraph(); p.add_run('A. 界面语言切换器（Globe 图标）'); style_heading(p, 3)
    add_bullet(doc, '**中文**：界面显示中文')
    add_bullet(doc, '**目标语言**：界面切换为当前学习的目标语言（如选英语词库，界面就变英文）')
    p = doc.add_paragraph(); p.add_run('B. 学习目标语言切换器（Languages 图标）'); style_heading(p, 3)
    add_para(doc, '选择你要学习的词库语言：')
    add_bullet(doc, '全部语言（综合复习）')
    add_bullet(doc, 'English / Japanese / Spanish / French / Portuguese')
    add_para(doc, '切换后，词库、统计、待复习词数都会立即按新语言重新加载。')

    p = doc.add_paragraph(); p.add_run('3.3 虚拟时间时钟（桌面端）'); style_heading(p, 2)
    add_para(doc, '导航栏右侧显示当前系统时间（含星期）。如果你触发了"时光旅行"，这里会显示虚拟时间。')

    doc.add_page_break()

    # ====================== 4. 仪表盘 ======================
    p = doc.add_paragraph(); p.add_run('4. 仪表盘（首页）'); style_heading(p, 1)
    add_para(doc, '仪表盘是你每次登录后的默认页面，集中查看学习状态。')

    p = doc.add_paragraph(); p.add_run('4.1 欢迎横幅'); style_heading(p, 2)
    add_para(doc, '紫色的顶部横幅显示主操作按钮（根据状态自动切换）：')
    add_bullet(doc, '有待复习词 → **"开始今日复习（N）"**（跳转复习页）')
    add_bullet(doc, '词库为空 → **"去添加我的第一个单词"**（跳转词库页）')
    add_bullet(doc, '今日已复习完 → 显示徽章 **"今日全部复习完成 🏆"**')
    add_para(doc, '副按钮 **"记忆曲线说明"** 可展开算法规则帮助卡，解释间隔、快速晋级、重置惩罚三大规则。')

    p = doc.add_paragraph(); p.add_run('4.2 三张统计卡'); style_heading(p, 2)
    add_table_from_rows(
        doc,
        ['卡片', '含义'],
        [
            ['已学单词', '当前语言下词库总数'],
            ['今日待复习', '已到复习时间且未掌握的词数（数字不为 0 时图标会脉冲闪烁）'],
            ['已掌握', '进入终期掌握阶段的词数（阶段 ≥ 5 或连续 3 次首通正确）'],
        ],
        col_widths_cm=[4, 12]
    )

    p = doc.add_paragraph(); p.add_run('4.3 周复习趋势图'); style_heading(p, 2)
    add_para(doc, '显示最近 7 天的复习数据：')
    add_bullet(doc, '**柱状图**：绿色 = 首通正确，红色 = 首通错误（堆叠显示当日总复习量）')
    add_bullet(doc, '**折线图**：蓝色 = 当日首通正确率（0–100%）')
    add_bullet(doc, '**右上角汇总**：本周总复习次数、平均首通正确率')
    add_bullet(doc, '鼠标悬停柱子可看详细数据（日期 / 正确数 / 错误数 / 正确率）')

    p = doc.add_paragraph(); p.add_run('4.4 记忆阶段分布图'); style_heading(p, 2)
    add_para(doc, '7 行水平进度条，直观展示你的词库在 7 个阶段的分布：')
    add_table_from_rows(
        doc,
        ['阶段', '含义', '下次复习间隔'],
        [
            ['阶段 0', '新词加深', '+1 天'],
            ['阶段 1', '第二天复习', '+2 天'],
            ['阶段 2', '第四天巩固', '+4 天'],
            ['阶段 3', '第七天强化', '+7 天'],
            ['阶段 4', '第 15 天熟记', '+15 天'],
            ['阶段 5', '第 30 天定型', '+30 天'],
            ['阶段 6（掌握）', '终期完全掌握', '不再排期'],
        ],
        col_widths_cm=[4, 6, 6]
    )
    add_tip(doc, '进阶提示：连续 3 次首轮答对可瞬间跃迁到阶段 5，大幅缩短掌握周期。', kind='tip')

    doc.add_page_break()

    # ====================== 5. 词库管理 ======================
    p = doc.add_paragraph(); p.add_run('5. 词库管理'); style_heading(p, 1)
    add_para(doc, '点击导航栏 **"词库"** 进入。这是你管理单词的核心页面。')

    p = doc.add_paragraph(); p.add_run('5.1 顶部统计概览'); style_heading(p, 2)
    add_para(doc, '显示当前语言下的整体掌握率（绿色进度条）、4 个指标卡（词库容量 / 待复习 / 进行中 / 已掌握）、7 阶段彩色柱状分布图。')

    p = doc.add_paragraph(); p.add_run('5.2 添加单词'); style_heading(p, 2)
    add_para(doc, '左侧面板提供两种添加方式，通过顶部 Tab 切换。')

    p = doc.add_paragraph(); p.add_run('方式一：单个录入'); style_heading(p, 3)
    add_numbered(doc, '在输入框输入一个单词（如 `ephemeral`）。')
    add_numbered(doc, '点击 **"添加"** 按钮。')
    add_numbered(doc, '系统会调用 AI 自动生成：')
    add_bullet(doc, '中文释义')
    add_bullet(doc, '双语例句（含中文翻译）')
    add_bullet(doc, '艾宾浩斯助记（词根 / 联想 / 谐音等记忆法）')
    add_bullet(doc, '音标、真人发音 URL')
    add_numbered(doc, '成功后顶部出现绿色提示，输入框清空，可继续添加下一个。')
    add_tip(doc, '若提示 "单词已存在"，说明该词已在你的词库中，不会重复添加。', kind='warn')

    p = doc.add_paragraph(); p.add_run('方式二：批量导入'); style_heading(p, 3)
    add_para(doc, '支持两种输入方式：')
    add_bullet(doc, '**拖拽 / 点击上传 .txt 文件**：点击虚线区域或直接把 txt 文件拖进去')
    add_bullet(doc, '**直接粘贴文本**：每行一个词，或用逗号、分号、空格分隔')
    add_numbered(doc, '输入内容后，文本框上方会显示 **"检测到 N 个"**。')
    add_numbered(doc, '点击 **"开始导入"**。')
    add_numbered(doc, '导入过程中显示进度条与当前处理的单词（顺序执行，非并发）。')
    add_numbered(doc, '完成后：绿色列出新增词，琥珀色列出重复或失败的词及原因。')
    add_tip(doc, '建议单次批量导入控制在 30–50 个以内，避免 AI 接口超时。', kind='tip')

    p = doc.add_paragraph(); p.add_run('5.3 浏览与筛选词汇'); style_heading(p, 2)
    add_para(doc, '右侧列表区支持多种筛选：')
    add_bullet(doc, '**状态 Tab**：全部 / 待复习 / 进行中 / 已掌握（含计数）')
    add_bullet(doc, '**阶段筛选**：下拉选择阶段 0–6')
    add_bullet(doc, '**排序**：最新 / 最旧 / 复习临近 / 字母顺序 / 阶段顺序')
    add_bullet(doc, '**每页数量**：10 / 15 / 20 / 30 / 50')
    add_bullet(doc, '**搜索框**：可匹配拼写、释义、例句、助记（不区分大小写）')

    p = doc.add_paragraph(); p.add_run('5.4 单词详情弹窗'); style_heading(p, 2)
    add_para(doc, '点击任意词条打开详情弹窗，包含：')
    p = doc.add_paragraph(); p.add_run('头部'); style_heading(p, 3)
    add_bullet(doc, '单词拼写（大字）')
    add_bullet(doc, '🔊 **发音按钮**：优先播放真人音频，失败时回退到浏览器语音合成')
    add_bullet(doc, '音标')
    p = doc.add_paragraph(); p.add_run('正文（4 个字段）'); style_heading(p, 3)
    add_bullet(doc, '**释义**')
    add_bullet(doc, '**双语例句**（带边框块）')
    add_bullet(doc, '**艾宾浩斯助记**（琥珀色块，带 Sparkles 图标）')
    add_bullet(doc, '**学习参数**（4 格）：学习阶段 / 连胜记录（X/3）/ 基准日期 / 下次复习日期')
    p = doc.add_paragraph(); p.add_run('底部操作'); style_heading(p, 3)
    add_bullet(doc, '🗑️ **删除**：彻底删除该词（会弹窗二次确认，操作不可逆）')
    add_bullet(doc, '✨ **AI 智能重新润色**：重新生成释义、例句、助记')
    add_bullet(doc, '✏️ **编辑详情**：切换到编辑模式，可手动修改音标、释义、例句、翻译、助记')
    add_bullet(doc, '💾 **保存修改** / **取消**（编辑模式下显示）')

    doc.add_page_break()

    # ====================== 6. 复习会话 ======================
    p = doc.add_paragraph(); p.add_run('6. 复习会话'); style_heading(p, 1)

    p = doc.add_paragraph(); p.add_run('6.1 进入复习'); style_heading(p, 2)
    add_para(doc, '有两种入口：')
    add_bullet(doc, '仪表盘点击 **"开始今日复习（N）"**')
    add_bullet(doc, '导航栏点击 **"复习"**（有待复习词时右上角显示红色徽章数字）')
    add_para(doc, '若没有待复习词，会显示空态卡：**"暂无待复习单词"**。')

    p = doc.add_paragraph(); p.add_run('6.2 选择学习模式'); style_heading(p, 2)
    add_para(doc, '开始会话前会让你二选一：')
    add_table_from_rows(
        doc,
        ['模式', '图标', '适用场景'],
        [
            ['闪卡模式', 'FileText', '卡片翻转记忆，适合快速自测与浏览释义'],
            ['拼写测试', 'Compass', '键入拼写核对，深度建立字母与音节肌肉记忆'],
        ],
        col_widths_cm=[4, 3, 9]
    )
    add_para(doc, '选择后点击 **"启动会话"**。')

    p = doc.add_paragraph(); p.add_run('6.3 闪卡模式'); style_heading(p, 2)
    add_para(doc, '**界面元素**：')
    add_bullet(doc, '**顶部**：当前轮次徽章、错词重考数量徽章、进度（X/Y）')
    add_bullet(doc, '**中央卡**：点击卡片任意区域可翻转')
    add_bullet(doc, '背面：拼写 + 音标 + 🔊 发音按钮 + **"显示翻译"** 按钮')
    add_bullet(doc, '正面：释义 + 双语例句 + 助记')
    add_bullet(doc, '**底部决策按钮**（仅在翻面后出现）：')
    add_bullet(doc, '🟢 **"记得"**（绿色）—— 记对了', level=1)
    add_bullet(doc, '🔴 **"忘了"**（红色）—— 记错了', level=1)
    add_tip(doc, '每张卡片出现时会自动播放发音。', kind='info')

    p = doc.add_paragraph(); p.add_run('6.4 拼写测试模式'); style_heading(p, 2)
    add_para(doc, '**操作流程**：')
    add_numbered(doc, '屏幕显示：音标 + 释义 + 挖空例句（拼写位置显示为 `首字母 + 下划线`，如 `e________`）')
    add_numbered(doc, '在输入框键入你记忆中的拼写')
    add_numbered(doc, '点击 **"检查"**（或回车）')
    add_numbered(doc, '系统判定后：✅ 正确输入框变绿；❌ 错误输入框变红并显示正确拼写供你对照')
    add_numbered(doc, '点击 **"下一步"** 继续下一个词')
    add_tip(doc, '输入框已禁用浏览器自动补全和拼写检查，不会给你"提示"。', kind='info')

    p = doc.add_paragraph(); p.add_run('6.5 错词循环重考（核心机制）'); style_heading(p, 2)
    add_para(doc, '**规则**：')
    add_bullet(doc, '复习中**答错的词会被自动加入本会话的"重考队列"**')
    add_bullet(doc, '本轮所有词过完后，若重考队列非空，会**自动开启新一轮**，只考刚才错的词')
    add_bullet(doc, '错词在新一轮中答对后，即从重考队列移除')
    add_bullet(doc, '直到重考队列清空，会话才算真正结束')
    add_tip(doc, '重要：最终的"首通正确率"统计只看每词的第一次作答，重考答对不影响首通数据。', kind='warn')

    p = doc.add_paragraph(); p.add_run('6.6 会话完成'); style_heading(p, 2)
    add_para(doc, '完成所有复习后显示结算页：')
    add_bullet(doc, '🏆 大奖牌图标')
    add_bullet(doc, '**3 项核心数据**：总复习数 / 首通正确数 / 首通正确率（≥80% 高亮绿色）')
    add_bullet(doc, '**AI 评语**：根据正确率动态生成鼓励或建议')
    add_bullet(doc, '**首通错误词列表**（若有错词）：列出所有首通答错的词，便于回顾')
    add_bullet(doc, '点击 **"同步记录并结束复习"** 提交后端、返回仪表盘')

    doc.add_page_break()

    # ====================== 7. 个人资料与勋章墙 ======================
    p = doc.add_paragraph(); p.add_run('7. 个人资料与勋章墙'); style_heading(p, 1)
    add_para(doc, '点击导航栏 **"个人"**（你的昵称）进入。')

    p = doc.add_paragraph(); p.add_run('7.1 账号横幅'); style_heading(p, 2)
    add_para(doc, '显示：头像（昵称前 2 字符）、用户名、邮箱、注册日期。右侧有 **"安全退出登录"** 按钮。')

    p = doc.add_paragraph(); p.add_run('7.2 修改基本资料'); style_heading(p, 2)
    add_bullet(doc, '**账号绑定邮箱**：不可更改（灰显）')
    add_bullet(doc, '**个性昵称**：可修改')
    add_bullet(doc, '**每日复习目标**：可修改（5–200）')
    add_para(doc, '点击 **"保存基本资料"** 同步到云端。')

    p = doc.add_paragraph(); p.add_run('7.3 修改登录密码'); style_heading(p, 2)
    add_para(doc, '填写旧密码、新密码（≥6 位）、确认新密码后点击 **"修改登录密码"**。前端会校验两次新密码是否一致。')

    p = doc.add_paragraph(); p.add_run('7.4 勋章墙（成就系统）'); style_heading(p, 2)
    add_para(doc, '顶部显示：已点亮勋章数（共 12 枚）、当前连续复习天数、历史最高连续天数。')

    p = doc.add_paragraph(); p.add_run('🏆 学习坚持（5 枚，按历史最高连续天数解锁）'); style_heading(p, 3)
    add_table_from_rows(
        doc,
        ['勋章', '解锁条件'],
        [
            ['初试锋芒', '连续复习 1 天'],
            ['渐入佳境', '连续复习 3 天'],
            ['习惯成自然', '连续复习 7 天'],
            ['意志钢铁', '连续复习 15 天'],
            ['艾宾极境之王', '连续复习 30 天'],
        ],
        col_widths_cm=[5, 11]
    )

    p = doc.add_paragraph(); p.add_run('📚 词库探索（4 枚，按词库累计词数解锁）'); style_heading(p, 3)
    add_table_from_rows(
        doc,
        ['勋章', '解锁条件'],
        [
            ['词汇起航', '词库累计 5 词'],
            ['积沙成塔', '词库累计 20 词'],
            ['学富五车', '词库累计 50 词'],
            ['词海逐浪', '词库累计 100 词'],
        ],
        col_widths_cm=[5, 11]
    )

    p = doc.add_paragraph(); p.add_run('⭐ 终极掌握（3 枚，按已掌握词数解锁）'); style_heading(p, 3)
    add_table_from_rows(
        doc,
        ['勋章', '解锁条件'],
        [
            ['初显成效', '掌握 3 词'],
            ['大显身手', '掌握 10 词'],
            ['融会贯通', '掌握 30 词'],
        ],
        col_widths_cm=[5, 11]
    )

    p = doc.add_paragraph(); p.add_run('勋章状态与详情'); style_heading(p, 3)
    add_bullet(doc, '**已解锁**：彩色渐变图标 + 光晕效果 + "已点亮"水印')
    add_bullet(doc, '**未解锁**：灰色 + 锁图标')
    add_bullet(doc, '每枚勋章都显示**当前进度条**（如 17/30 天）')
    add_bullet(doc, '点击任意勋章卡可查看详情弹窗：完整描述、解锁寄语（已解锁）、当前进度')
    add_para(doc, '可用顶部 Tab 按类别筛选：全部 / 学习坚持 / 词库探索 / 终极掌握。')

    doc.add_page_break()

    # ====================== 8. 记忆算法原理 ======================
    p = doc.add_paragraph(); p.add_run('8. 记忆算法原理'); style_heading(p, 1)
    add_para(doc, '理解算法能帮助你更高效地学习。')

    p = doc.add_paragraph(); p.add_run('8.1 6 阶段间隔重复'); style_heading(p, 2)
    add_para(doc, '单词从新词到完全掌握，需要经过 6 个阶段，每阶段复习间隔如下：')
    add_code_block(doc, '新词 → 阶段 0 → 阶段 1 → 阶段 2 → 阶段 3 → 阶段 4 → 阶段 5 → 已掌握\n       +1天      +2天      +4天     +7天     +15天    +30天')

    p = doc.add_paragraph(); p.add_run('8.2 答对的奖励'); style_heading(p, 2)
    add_para(doc, '每次**首轮答对**（first try correct）：')
    add_numbered(doc, '连胜数 +1')
    add_numbered(doc, '**若连胜数 ≥ 3**：直接跃迁到阶段 5（下次复习 = 30 天后），进入快速掌握通道')
    add_numbered(doc, '**否则**：正常升一阶，按新阶段的间隔排下次复习')

    p = doc.add_paragraph(); p.add_run('8.3 答错的惩罚'); style_heading(p, 2)
    add_para(doc, '**首轮答错**会触发重置：')
    add_numbered(doc, '阶段回到 0')
    add_numbered(doc, '连胜数清零')
    add_numbered(doc, '下次复习 = 明天（即第二天立即重练）')

    p = doc.add_paragraph(); p.add_run('8.4 "已掌握"的判定'); style_heading(p, 2)
    add_para(doc, '满足以下任一条件即视为已掌握：')
    add_bullet(doc, '阶段 ≥ 5（通过完整间隔升上来）')
    add_bullet(doc, '或连胜数 ≥ 3（通过快速通道）')
    add_para(doc, '已掌握的词不再排期复习，但你可以到词库手动管理它们。')

    p = doc.add_paragraph(); p.add_run('8.5 错词循环 vs 首通正确率'); style_heading(p, 2)
    add_bullet(doc, '**错词循环**：会话内立即重考，确保你在本次复习中"记住"了它')
    add_bullet(doc, '**首通正确率**：只统计每词的第一次作答结果，反映真实记忆水平')
    add_para(doc, '两者配合：错词循环保证短期记忆形成，首通正确率反映长期学习成效。')

    doc.add_page_break()

    # ====================== 9. 管理员专属功能 ======================
    p = doc.add_paragraph(); p.add_run('9. 管理员专属功能'); style_heading(p, 1)
    add_tip(doc, '以下两个功能仅对管理员邮箱 wujizong@gmail.com 可见，其他普通用户在仪表盘上看不到这两块卡片，记忆阶段分布图会自动扩展为整行显示。', kind='danger')

    p = doc.add_paragraph(); p.add_run('9.1 SRS 时光机'); style_heading(p, 2)
    add_para(doc, '用于**演示和调试**场景，可以快进系统时间，快速触发复习。')
    add_table_from_rows(
        doc,
        ['按钮', '效果'],
        [
            ['+1 天', '时间快进 1 天'],
            ['+3 天', '时间快进 3 天'],
            ['+7 天', '时间快进 7 天'],
            ['+30 天', '时间快进 30 天'],
        ],
        col_widths_cm=[5, 11]
    )
    add_para(doc, '快进后：')
    add_bullet(doc, '仪表盘出现琥珀色提示条"已快进 N 天"')
    add_bullet(doc, '页面底部出现全局提示条')
    add_bullet(doc, '导航栏时钟显示虚拟时间')
    add_bullet(doc, '所有"待复习词数"按虚拟时间重新计算')
    add_para(doc, '点击 **"还原今天"** 可回到真实时间。')

    p = doc.add_paragraph(); p.add_run('9.2 重置系统数据库'); style_heading(p, 2)
    add_tip(doc, '高危操作：会清除当前语言下的所有数据并重置为初始状态。', kind='danger')
    add_para(doc, '点击 **"重置系统数据库"** 会弹出红色警告和二次确认：')
    add_bullet(doc, '**"确定重置"**：执行重置（不可恢复）')
    add_bullet(doc, '**"取消"**：放弃操作')

    doc.add_page_break()

    # ====================== 10. FAQ ======================
    p = doc.add_paragraph(); p.add_run('10. 常见问题（FAQ）'); style_heading(p, 1)

    faqs = [
        ('Q1：添加单词后 AI 释义生成失败怎么办？',
         '通常是 AI 服务临时不可用。你可以：1. 稍后重试添加；2. 词库中点击该词详情，使用 **"AI 智能重新润色"** 重新生成；3. 也可以手动编辑详情填入自己的释义。'),
        ('Q2：点击发音按钮没有声音？',
         '可能原因：浏览器未允许自动播放音频 → 点击页面任意位置后再试；系统音量为 0 或静音；网络问题导致音频 URL 加载失败 → 系统会自动回退到浏览器 TTS 语音合成。'),
        ('Q3：批量导入卡住不动？',
         '批量导入是顺序执行（避免并发触发 AI 限流），导入大量单词时需要耐心等待。建议单次控制在 30–50 个以内。'),
        ('Q4：忘记密码怎么办？',
         '应用暂未提供邮箱找回密码。请联系管理员协助重置，或通过其他已登录设备到个人资料页修改。'),
        ('Q5：连续复习天数是怎么计算的？',
         '系统按你每天是否有复习记录来计算连续天数。某天没有复习记录，连续天数会中断。建议每天至少完成一次复习会话以保持连胜。'),
        ('Q6：切换目标语言后，之前的词库会丢失吗？',
         '不会。所有词库都按语言分类存储，切换语言只是筛选显示。切回原来的语言即可看到之前的所有单词。'),
        ('Q7：已掌握的词还能再复习吗？',
         '系统默认不再为已掌握的词排期。但你可以到词库中点击该词查看详情，或删除后重新添加来重新学习。'),
        ('Q8：管理员功能（时光机 / 数据库重置）普通用户为什么看不到？',
         '这两个功能属于高危操作，仅对管理员邮箱开放。普通用户仪表盘不会显示这两块卡片，避免误操作带来数据风险。如果你确实需要这些功能，请联系管理员。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r = p.add_run(q)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = INDIGO
        r.font.name = EN_FONT
        rpr = r._element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), CN_FONT)
        rpr.append(rfonts)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

        p2 = doc.add_paragraph()
        add_runs_with_inline(p2, a)
        set_paragraph_cjk_font(p2, size=11, color=SLATE_700)
        p2.paragraph_format.line_spacing = 1.5
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(6)

    # 结语
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    add_tip(doc, '结语：艾宾浩斯记忆法的关键在于坚持。每天打开应用完成今日复习，连续 30 天就能解锁最高勋章"艾宾极境之王"。祝你学习愉快！', kind='tip')


def main():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认正文样式
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), CN_FONT)

    build_cover(doc)
    build_toc(doc)
    build_body(doc)

    doc.save(OUT)
    print('已生成: ' + OUT)
    print('大小: ' + str(os.path.getsize(OUT)) + ' bytes')


if __name__ == '__main__':
    main()
